from flask import Flask, request, jsonify, render_template, redirect, url_for, flash, send_file
from flask_jwt_extended import JWTManager, create_access_token, jwt_required, get_jwt_identity
from flask_bcrypt import Bcrypt
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
import json
import os
from datetime import datetime, timedelta
import sqlite3
from threading import Lock
import logging
import io
from io import BytesIO

# Импорты для генерации отчетов
import pandas as pd
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Inches
import tempfile

from config import Config

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Конфигурация Flask
app.config['SECRET_KEY'] = Config.SECRET_KEY
app.config['JWT_SECRET_KEY'] = Config.JWT_SECRET_KEY
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=Config.JWT_ACCESS_TOKEN_EXPIRES_HOURS)

# Инициализация расширений
jwt = JWTManager(app)
bcrypt = Bcrypt(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'web_login'
login_manager.login_message = 'Пожалуйста, войдите для доступа к этой странице.'

# Блокировка для потокобезопасности БД
db_lock = Lock()

# Папки для данных
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, Config.DATA_DIR)
DB_PATH = os.path.join(DATA_DIR, Config.DB_NAME)
os.makedirs(DATA_DIR, exist_ok=True)

# Инициализация Fernet для шифрования
fernet = Config.get_fernet()


# Модель пользователя для Flask-Login
class User(UserMixin):
    def __init__(self, id, username, is_admin=False):
        self.id = id
        self.username = username
        self.is_admin = bool(is_admin)


def get_db_connection():
    """Создание подключения к базе данных"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


@login_manager.user_loader
def load_user(user_id):
    """Загрузка пользователя для Flask-Login"""
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM users WHERE id = ?', (user_id,))
    user_data = cursor.fetchone()
    conn.close()

    if user_data:
        is_admin = user_data['is_admin'] if 'is_admin' in user_data.keys() else False
        return User(user_data['id'], user_data['username'], is_admin)
    return None


def check_and_update_database():
    """Проверка и обновление структуры базы данных"""
    logger.info("Checking database structure...")

    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()

    try:
        # Получаем информацию о таблице devices
        cursor.execute("PRAGMA table_info(devices)")
        columns = [column[1] for column in cursor.fetchall()]
        logger.info(f"Existing columns in devices table: {columns}")

        # Список новых колонок для добавления
        new_columns = [
            ('direction', 'TEXT'),
            ('inventory_number', 'TEXT'),
            ('room', 'TEXT'),
            ('notes', 'TEXT'),
            ('is_deleted', 'BOOLEAN DEFAULT 0'),
            ('deleted_at', 'DATETIME')
        ]

        # Добавляем отсутствующие колонки
        for column_name, column_type in new_columns:
            if column_name not in columns:
                logger.info(f"Adding column {column_name} to devices table")
                try:
                    cursor.execute(f'ALTER TABLE devices ADD COLUMN {column_name} {column_type}')
                    logger.info(f"Column {column_name} added successfully")
                except Exception as e:
                    logger.error(f"Error adding column {column_name}: {e}")
                    continue

        # Создаем таблицу для логов действий
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS action_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                username TEXT,
                action TEXT,
                details TEXT,
                ip_address TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')

        conn.commit()
        logger.info("Database structure update completed")

    except Exception as e:
        logger.error(f"Error updating database structure: {e}")
        conn.rollback()
    finally:
        conn.close()


def init_database():
    """Инициализация базы данных"""
    logger.info(f"Initializing database at: {DB_PATH}")

    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            # Таблица устройств (базовая структура)
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS devices (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    device_id TEXT UNIQUE,
                    computer_name TEXT,
                    mac_address TEXT,
                    cpu_info TEXT,
                    gpu_info TEXT,
                    memory_info TEXT,
                    disk_info TEXT,
                    os_info TEXT,
                    architecture TEXT,
                    python_version TEXT,
                    ip_address TEXT,
                    timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                    last_updated DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            # Таблица пользователей для авторизации
            cursor.execute('''
                CREATE TABLE IF NOT EXISTS users (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    username TEXT UNIQUE NOT NULL,
                    password_hash TEXT NOT NULL,
                    is_admin BOOLEAN DEFAULT 0,
                    is_active BOOLEAN DEFAULT 1,
                    created_at DATETIME DEFAULT CURRENT_TIMESTAMP
                )
            ''')

            # Создаем администратора по умолчанию
            admin_password = bcrypt.generate_password_hash('admin123').decode('utf-8')
            cursor.execute('''
                INSERT OR IGNORE INTO users (username, password_hash, is_admin) 
                VALUES (?, ?, ?)
            ''', ('admin', admin_password, 1))

            conn.commit()
            logger.info("Database tables created successfully")

        except Exception as e:
            logger.error(f"Error creating database tables: {e}")
            conn.rollback()
            raise
        finally:
            conn.close()

    # Запускаем миграцию отдельно после создания базовых таблиц
    check_and_update_database()


def save_device_data(data):
    """Сохранение или обновление данных устройства"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            # Проверяем существование устройства
            cursor.execute('SELECT id FROM devices WHERE device_id = ?', (data['device_id'],))
            existing = cursor.fetchone()

            if existing:
                # Обновляем существующую запись
                cursor.execute('''
                    UPDATE devices SET 
                    computer_name = ?, mac_address = ?, cpu_info = ?, gpu_info = ?,
                    memory_info = ?, disk_info = ?, os_info = ?, architecture = ?,
                    python_version = ?, ip_address = ?, last_updated = CURRENT_TIMESTAMP,
                    direction = ?, inventory_number = ?, room = ?, notes = ?
                    WHERE device_id = ?
                ''', (
                    data['computer_name'], data['mac_address'], data['cpu_info'],
                    data['gpu_info'], data['memory_info'], data['disk_info'],
                    data['os_info'], data['architecture'], data['python_version'],
                    request.remote_addr,
                    data.get('direction', ''),
                    data.get('inventory_number', ''),
                    data.get('room', ''),
                    data.get('notes', ''),
                    data['device_id']
                ))
            else:
                # Вставляем новую запись
                cursor.execute('''
                    INSERT INTO devices 
                    (device_id, computer_name, mac_address, cpu_info, gpu_info,
                     memory_info, disk_info, os_info, architecture, python_version, ip_address,
                     direction, inventory_number, room, notes)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (
                    data['device_id'], data['computer_name'], data['mac_address'],
                    data['cpu_info'], data['gpu_info'], data['memory_info'],
                    data['disk_info'], data['os_info'], data['architecture'],
                    data['python_version'], request.remote_addr,
                    data.get('direction', ''),
                    data.get('inventory_number', ''),
                    data.get('room', ''),
                    data.get('notes', '')
                ))

            conn.commit()
            logger.info(f"Data saved for device: {data['device_id']}")
            return True

        except Exception as e:
            logger.error(f"Error saving device data: {e}")
            conn.rollback()
            return False
        finally:
            conn.close()


def get_all_devices(include_deleted=False):
    """Получение всех устройств"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            if include_deleted:
                cursor.execute('''
                    SELECT * FROM devices 
                    ORDER BY last_updated DESC
                ''')
            else:
                cursor.execute('''
                    SELECT * FROM devices 
                    WHERE is_deleted = 0 OR is_deleted IS NULL
                    ORDER BY last_updated DESC
                ''')
            devices = []
            for row in cursor.fetchall():
                device_dict = dict(row)
                for field in ['direction', 'inventory_number', 'room', 'notes']:
                    if field not in device_dict or device_dict[field] is None:
                        device_dict[field] = ''
                devices.append(device_dict)
            return devices
        except Exception as e:
            logger.error(f"Error getting devices: {e}")
            return []
        finally:
            conn.close()


def get_deleted_devices():
    """Получение удаленных устройств"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('''
                SELECT * FROM devices 
                WHERE is_deleted = 1
                ORDER BY deleted_at DESC
            ''')
            devices = []
            for row in cursor.fetchall():
                device_dict = dict(row)
                for field in ['direction', 'inventory_number', 'room', 'notes']:
                    if field not in device_dict or device_dict[field] is None:
                        device_dict[field] = ''
                devices.append(device_dict)
            return devices
        except Exception as e:
            logger.error(f"Error getting deleted devices: {e}")
            return []
        finally:
            conn.close()


def get_device(device_id):
    """Получение конкретного устройства"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('SELECT * FROM devices WHERE device_id = ?', (device_id,))
            device = cursor.fetchone()
            if device:
                device_dict = dict(device)
                for field in ['direction', 'inventory_number', 'room', 'notes']:
                    if field not in device_dict or device_dict[field] is None:
                        device_dict[field] = ''
                return device_dict
            return None
        except Exception as e:
            logger.error(f"Error getting device {device_id}: {e}")
            return None
        finally:
            conn.close()


def get_all_users():
    """Получение всех пользователей"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('SELECT id, username, is_admin, is_active, created_at FROM users')
            users = []
            for row in cursor.fetchall():
                users.append(dict(row))
            return users
        except Exception as e:
            logger.error(f"Error getting users: {e}")
            return []
        finally:
            conn.close()


def create_user(username, password, is_admin=False):
    """Создание нового пользователя"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            password_hash = bcrypt.generate_password_hash(password).decode('utf-8')
            cursor.execute('''
                INSERT INTO users (username, password_hash, is_admin)
                VALUES (?, ?, ?)
            ''', (username, password_hash, is_admin))
            conn.commit()
            return True
        except sqlite3.IntegrityError:
            return False
        except Exception as e:
            logger.error(f"Error creating user: {e}")
            return False
        finally:
            conn.close()


def update_user_password(user_id, new_password):
    """Обновление пароля пользователя"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            # Получаем информацию о пользователе для логирования
            cursor.execute('SELECT username FROM users WHERE id = ?', (user_id,))
            user_data = cursor.fetchone()
            username = user_data['username'] if user_data else 'Unknown'

            password_hash = bcrypt.generate_password_hash(new_password).decode('utf-8')
            cursor.execute('''
                UPDATE users SET password_hash = ? WHERE id = ?
            ''', (password_hash, user_id))
            conn.commit()

            logger.info(f"Password updated for user: {username} (ID: {user_id})")
            return True
        except Exception as e:
            logger.error(f"Error updating user password for ID {user_id}: {e}")
            return False
        finally:
            conn.close()


def toggle_user_status(user_id):
    """Блокировка/разблокировка пользователя"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('SELECT is_active FROM users WHERE id = ?', (user_id,))
            current_status = cursor.fetchone()
            if current_status:
                new_status = not current_status['is_active']
                cursor.execute('''
                    UPDATE users SET is_active = ? WHERE id = ?
                ''', (new_status, user_id))
                conn.commit()
                return True
            return False
        except Exception as e:
            logger.error(f"Error toggling user status: {e}")
            return False
        finally:
            conn.close()


def soft_delete_device(device_id):
    """Мягкое удаление устройства"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('''
                UPDATE devices SET is_deleted = 1, deleted_at = CURRENT_TIMESTAMP
                WHERE device_id = ?
            ''', (device_id,))
            conn.commit()
            return cursor.rowcount > 0
        except Exception as e:
            logger.error(f"Error soft deleting device: {e}")
            return False
        finally:
            conn.close()


def restore_device(device_id):
    """Восстановление устройства"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('''
                UPDATE devices SET is_deleted = 0, deleted_at = NULL
                WHERE device_id = ?
            ''', (device_id,))
            conn.commit()
            return cursor.rowcount > 0
        except Exception as e:
            logger.error(f"Error restoring device: {e}")
            return False
        finally:
            conn.close()


def get_duplicate_inventory_numbers():
    """Поиск дублирующихся инвентарных номеров"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('''
                SELECT inventory_number, COUNT(*) as count
                FROM devices 
                WHERE inventory_number IS NOT NULL 
                AND inventory_number != '' 
                AND (is_deleted = 0 OR is_deleted IS NULL)
                GROUP BY inventory_number 
                HAVING COUNT(*) > 1
            ''')
            duplicates = {}
            for row in cursor.fetchall():
                duplicates[row['inventory_number']] = row['count']
            return duplicates
        except Exception as e:
            logger.error(f"Error getting duplicate inventory numbers: {e}")
            return {}
        finally:
            conn.close()


def get_unique_directions():
    """Получение уникальных направлений"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('''
                SELECT DISTINCT direction 
                FROM devices 
                WHERE direction IS NOT NULL 
                AND direction != '' 
                AND (is_deleted = 0 OR is_deleted IS NULL)
                ORDER BY direction
            ''')
            directions = [row['direction'] for row in cursor.fetchall()]
            return directions
        except Exception as e:
            logger.error(f"Error getting unique directions: {e}")
            return []
        finally:
            conn.close()


def authenticate_user(username, password):
    """Аутентификация пользователя"""
    conn = get_db_connection()
    cursor = conn.cursor()

    try:
        cursor.execute('SELECT * FROM users WHERE username = ? AND is_active = 1', (username,))
        user_row = cursor.fetchone()

        if user_row and bcrypt.check_password_hash(user_row['password_hash'], password):
            is_admin = user_row['is_admin'] if 'is_admin' in user_row.keys() else False
            return User(user_row['id'], user_row['username'], is_admin)
        return None
    except Exception as e:
        logger.error(f"Error authenticating user {username}: {e}")
        return None
    finally:
        conn.close()


def log_action(user_id, username, action, details):
    """Логирование действий пользователей"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()

        cursor.execute('''
            INSERT INTO action_logs (user_id, username, action, details, ip_address)
            VALUES (?, ?, ?, ?, ?)
        ''', (user_id, username, action, details, request.remote_addr))

        conn.commit()
        conn.close()

        logger.info(f"Action logged: {username} - {action} - {details}")
    except Exception as e:
        logger.error(f"Error logging action: {e}")


def get_action_logs(limit=100):
    """Получение логов действий"""
    with db_lock:
        conn = get_db_connection()
        cursor = conn.cursor()

        try:
            cursor.execute('''
                SELECT * FROM action_logs 
                ORDER BY timestamp DESC 
                LIMIT ?
            ''', (limit,))
            logs = []
            for row in cursor.fetchall():
                logs.append(dict(row))
            return logs
        except Exception as e:
            logger.error(f"Error getting action logs: {e}")
            return []
        finally:
            conn.close()


def encrypt_data(data):
    """Шифрование данных"""
    try:
        if isinstance(data, str):
            data = data.encode('utf-8')
        encrypted = fernet.encrypt(data)
        return encrypted.decode('utf-8')
    except Exception as e:
        logger.error(f"Encryption error: {e}")
        raise


def decrypt_data(encrypted_data):
    """Дешифрование данных"""
    try:
        if isinstance(encrypted_data, str):
            encrypted_data = encrypted_data.encode('utf-8')
        decrypted = fernet.decrypt(encrypted_data)
        return decrypted.decode('utf-8')
    except Exception as e:
        logger.error(f"Decryption error: {e}")
        raise


# ==================== ФУНКЦИИ ГЕНЕРАЦИИ ОТЧЕТОВ ====================

def generate_excel_report(devices):
    """Генерация отчета в формате Excel"""
    try:
        # Создаем workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Устройства"

        # Заголовки
        headers = [
            'ID устройства', 'Имя компьютера', 'MAC адрес', 'Процессор',
            'Видеокарта', 'Память', 'Диски', 'ОС', 'Архитектура',
            'Python версия', 'Направление', 'Инвентарный номер',
            'Кабинет', 'Примечание', 'Последнее обновление'
        ]

        # Стили для заголовков
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")

        # Записываем заголовки
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

        # Записываем данные
        for row, device in enumerate(devices, 2):
            ws.cell(row=row, column=1, value=device.get('device_id', ''))
            ws.cell(row=row, column=2, value=device.get('computer_name', ''))
            ws.cell(row=row, column=3, value=device.get('mac_address', ''))
            ws.cell(row=row, column=4, value=device.get('cpu_info', ''))
            ws.cell(row=row, column=5, value=device.get('gpu_info', ''))
            ws.cell(row=row, column=6, value=device.get('memory_info', ''))
            ws.cell(row=row, column=7, value=device.get('disk_info', ''))
            ws.cell(row=row, column=8, value=device.get('os_info', ''))
            ws.cell(row=row, column=9, value=device.get('architecture', ''))
            ws.cell(row=row, column=10, value=device.get('python_version', ''))
            ws.cell(row=row, column=11, value=device.get('direction', ''))
            ws.cell(row=row, column=12, value=device.get('inventory_number', ''))
            ws.cell(row=row, column=13, value=device.get('room', ''))
            ws.cell(row=row, column=14, value=device.get('notes', ''))
            ws.cell(row=row, column=15, value=device.get('last_updated', ''))

        # Авто-ширина колонок
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min((max_length + 2), 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        # Сохраняем в bytes
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        logger.error(f"Error generating Excel report: {e}")
        raise


def generate_pdf_report(devices):
    """Генерация отчета в формате PDF"""
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        elements = []

        # Стили
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # Center
            textColor=colors.HexColor('#2c3e50')
        )

        # Заголовок
        title = Paragraph("ОТЧЕТ ПО УСТРОЙСТВАМ", title_style)
        elements.append(title)

        # Информация о генерации
        info_style = ParagraphStyle(
            'InfoStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.gray
        )
        info_text = f"Сгенерировано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Всего устройств: {len(devices)}"
        elements.append(Paragraph(info_text, info_style))
        elements.append(Spacer(1, 20))

        # Таблица с данными
        if devices:
            # Подготовка данных для таблицы
            table_data = [['ID устройства', 'Компьютер', 'Направление', 'Инв. номер', 'Кабинет']]

            for device in devices:
                table_data.append([
                    device.get('device_id', '')[:8],
                    device.get('computer_name', '')[:20],
                    device.get('direction', '')[:15],
                    device.get('inventory_number', '')[:10],
                    device.get('room', '')[:10]
                ])

            # Создание таблицы
            table = Table(table_data, colWidths=[1 * inch, 1.5 * inch, 1.2 * inch, 1 * inch, 0.8 * inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            elements.append(table)

        else:
            elements.append(Paragraph("Нет данных для отображения", styles['Normal']))

        doc.build(elements)
        buffer.seek(0)
        return buffer

    except Exception as e:
        logger.error(f"Error generating PDF report: {e}")
        raise


def generate_docx_report(devices):
    """Генерация отчета в формате Word"""
    try:
        doc = Document()

        # Заголовок
        title = doc.add_heading('ОТЧЕТ ПО УСТРОЙСТВАМ', 0)
        title.alignment = 1  # Center

        # Информация о генерации
        doc.add_paragraph(f"Сгенерировано: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Всего устройств: {len(devices)}")
        doc.add_paragraph()

        if devices:
            # Создаем таблицу
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            headers = ['ID устройства', 'Компьютер', 'Направление', 'Инв. номер', 'Кабинет']
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True

            # Данные устройств
            for device in devices:
                row_cells = table.add_row().cells
                row_cells[0].text = device.get('device_id', '')[:8]
                row_cells[1].text = device.get('computer_name', '')[:20]
                row_cells[2].text = device.get('direction', '')[:15]
                row_cells[3].text = device.get('inventory_number', '')[:10]
                row_cells[4].text = device.get('room', '')[:10]

        else:
            doc.add_paragraph('Нет данных для отображения')

        # Сохраняем в buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        logger.error(f"Error generating DOCX report: {e}")
        raise


# Инициализация базы данных при старте
try:
    logger.info("Starting database initialization...")
    init_database()
    logger.info("Database initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize database: {e}")
    logger.info("Continuing with basic functionality...")


# ==================== ВЕБ-МАРШРУТЫ С АВТОРИЗАЦИЕЙ ====================

@app.route('/')
@login_required
def index():
    """Главная страница панели управления"""
    try:
        devices_count = len(get_all_devices())
        deleted_devices_count = len(get_deleted_devices())
        users_count = len(get_all_users())

        # Логируем просмотр главной страницы
        log_action(current_user.id, current_user.username, 'view_page', 'Просмотр главной страницы')

        return render_template('dashboard.html',
                               devices_count=devices_count,
                               deleted_devices_count=deleted_devices_count,
                               users_count=users_count,
                               current_user=current_user)
    except Exception as e:
        logger.error(f"Error in index route: {e}")
        flash('Ошибка загрузки данных', 'error')
        # Передаем значения по умолчанию при ошибке
        return render_template('dashboard.html',
                               devices_count=0,
                               deleted_devices_count=0,
                               users_count=0,
                               current_user=current_user)


@app.route('/login', methods=['GET', 'POST'])
def web_login():
    """Страница входа в веб-интерфейс"""
    if current_user.is_authenticated:
        return redirect(url_for('index'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        remember = bool(request.form.get('remember'))

        user = authenticate_user(username, password)
        if user:
            login_user(user, remember=remember)

            # Логируем успешный вход
            log_action(user.id, user.username, 'login', 'Успешный вход в систему')

            next_page = request.args.get('next')
            flash(f'Добро пожаловать, {username}!', 'success')
            return redirect(next_page or url_for('index'))
        else:
            # Логируем неудачную попытку входа
            log_action(None, username, 'login_failed', 'Неудачная попытка входа')
            flash('Неверное имя пользователя или пароль', 'error')

    return render_template('login.html')


@app.route('/logout')
@login_required
def web_logout():
    """Выход из системы"""
    # Логируем выход
    log_action(current_user.id, current_user.username, 'logout', 'Выход из системы')

    logout_user()
    flash('Вы вышли из системы', 'info')
    return redirect(url_for('web_login'))


@app.route('/devices')
@login_required
def web_devices():
    """Страница со списком устройств"""
    try:
        direction_filter = request.args.get('direction', '')
        search_query = request.args.get('search', '')

        devices = get_all_devices()
        directions = get_unique_directions()
        duplicates = get_duplicate_inventory_numbers()

        # Применяем фильтры
        if direction_filter:
            devices = [d for d in devices if d.get('direction') == direction_filter]

        if search_query:
            search_lower = search_query.lower()
            devices = [d for d in devices if
                       search_lower in d.get('computer_name', '').lower() or
                       search_lower in d.get('inventory_number', '').lower() or
                       search_lower in d.get('device_id', '').lower()]

        # Логируем просмотр устройств
        log_action(current_user.id, current_user.username, 'view_page',
                   f'Просмотр списка устройств (фильтр: {direction_filter or "нет"})')

        return render_template('devices.html',
                               devices=devices,
                               directions=directions or [],
                               duplicates=duplicates or {},
                               current_direction=direction_filter,
                               search_query=search_query,
                               current_user=current_user)
    except Exception as e:
        logger.error(f"Error in web_devices route: {e}")
        flash('Ошибка загрузки устройств', 'error')
        return render_template('devices.html',
                               devices=[],
                               directions=[],
                               duplicates={},
                               current_direction='',
                               search_query='',
                               current_user=current_user)


@app.route('/device/<device_id>')
@login_required
def device_detail(device_id):
    """Страница с детальной информацией об устройстве"""
    try:
        device = get_device(device_id)
        if not device:
            flash('Устройство не найдено', 'error')
            return redirect(url_for('web_devices'))

        # Логируем просмотр устройства
        log_action(current_user.id, current_user.username, 'view_device',
                   f'Просмотр устройства {device_id}')

        return render_template('device_detail.html', device=device, current_user=current_user)
    except Exception as e:
        logger.error(f"Error in device_detail route: {e}")
        flash('Ошибка загрузки устройства', 'error')
        return redirect(url_for('web_devices'))


@app.route('/devices/deleted')
@login_required
def deleted_devices():
    """Страница с удаленными устройствами"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('index'))

    try:
        deleted_devices = get_deleted_devices()

        # Логируем просмотр удаленных устройств
        log_action(current_user.id, current_user.username, 'view_page', 'Просмотр корзины устройств')

        return render_template('deleted_devices.html',
                               devices=deleted_devices,
                               current_user=current_user)
    except Exception as e:
        logger.error(f"Error in deleted_devices route: {e}")
        flash('Ошибка загрузки удаленных устройств', 'error')
        return render_template('deleted_devices.html', devices=[], current_user=current_user)


@app.route('/users')
@login_required
def web_users():
    """Страница управления пользователями (только для админов)"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('index'))

    try:
        users = get_all_users()

        # Логируем просмотр пользователей
        log_action(current_user.id, current_user.username, 'view_page', 'Просмотр списка пользователей')

        return render_template('users.html', users=users, current_user=current_user)
    except Exception as e:
        logger.error(f"Error in web_users route: {e}")
        flash('Ошибка загрузки пользователей', 'error')
        return render_template('users.html', users=[], current_user=current_user)


@app.route('/users/create', methods=['POST'])
@login_required
def create_user_route():
    """Создание нового пользователя"""
    if not current_user.is_admin:
        return jsonify({'error': 'Доступ запрещен'}), 403

    try:
        username = request.form.get('username')
        password = request.form.get('password')
        is_admin = bool(request.form.get('is_admin'))

        if not username or not password:
            flash('Заполните все обязательные поля', 'error')
            return redirect(url_for('web_users'))

        if create_user(username, password, is_admin):
            # Логируем создание пользователя
            log_action(current_user.id, current_user.username, 'create_user',
                       f'Создан пользователь {username} (админ: {is_admin})')
            flash(f'Пользователь {username} успешно создан', 'success')
        else:
            flash('Ошибка создания пользователя. Возможно, пользователь уже существует.', 'error')

        return redirect(url_for('web_users'))

    except Exception as e:
        logger.error(f"Error creating user: {e}")
        flash('Ошибка создания пользователя', 'error')
        return redirect(url_for('web_users'))


@app.route('/users/<int:user_id>/toggle', methods=['POST'])
@login_required
def toggle_user_route(user_id):
    """Блокировка/разблокировка пользователя"""
    if not current_user.is_admin:
        return jsonify({'error': 'Доступ запрещен'}), 403

    try:
        if toggle_user_status(user_id):
            # Логируем изменение статуса пользователя
            log_action(current_user.id, current_user.username, 'toggle_user',
                       f'Изменение статуса пользователя ID {user_id}')
            flash('Статус пользователя изменен', 'success')
        else:
            flash('Ошибка изменения статуса пользователя', 'error')
        return redirect(url_for('web_users'))
    except Exception as e:
        logger.error(f"Error toggling user: {e}")
        flash('Ошибка изменения статуса пользователя', 'error')
        return redirect(url_for('web_users'))


@app.route('/users/<int:user_id>/password', methods=['POST'])
@login_required
def change_password_route(user_id):
    """Смена пароля пользователя"""
    if not current_user.is_admin:
        return jsonify({'error': 'Доступ запрещен'}), 403

    try:
        new_password = request.form.get('new_password')
        if not new_password:
            flash('Введите новый пароль', 'error')
            return redirect(url_for('web_users'))

        if update_user_password(user_id, new_password):
            # Логируем смену пароля
            log_action(current_user.id, current_user.username, 'change_password',
                       f'Смена пароля пользователя ID {user_id}')
            flash('Пароль успешно изменен', 'success')
        else:
            flash('Ошибка изменения пароля', 'error')
        return redirect(url_for('web_users'))
    except Exception as e:
        logger.error(f"Error changing password: {e}")
        flash('Ошибка изменения пароля', 'error')
        return redirect(url_for('web_users'))


@app.route('/device/<device_id>/delete', methods=['POST'])
@login_required
def delete_device_route(device_id):
    """Удаление устройства"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('web_devices'))

    try:
        if soft_delete_device(device_id):
            # Логируем удаление устройства
            log_action(current_user.id, current_user.username, 'delete_device',
                       f'Удаление устройства {device_id}')
            flash('Устройство перемещено в корзину', 'success')
        else:
            flash('Ошибка удаления устройства', 'error')
        return redirect(url_for('web_devices'))
    except Exception as e:
        logger.error(f"Error deleting device: {e}")
        flash('Ошибка удаления устройства', 'error')
        return redirect(url_for('web_devices'))


@app.route('/device/<device_id>/restore', methods=['POST'])
@login_required
def restore_device_route(device_id):
    """Восстановление устройства"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('deleted_devices'))

    try:
        if restore_device(device_id):
            # Логируем восстановление устройства
            log_action(current_user.id, current_user.username, 'restore_device',
                       f'Восстановление устройства {device_id}')
            flash('Устройство восстановлено', 'success')
        else:
            flash('Ошибка восстановления устройства', 'error')
        return redirect(url_for('deleted_devices'))
    except Exception as e:
        logger.error(f"Error restoring device: {e}")
        flash('Ошибка восстановления устройства', 'error')
        return redirect(url_for('deleted_devices'))


@app.route('/settings')
@login_required
def web_settings():
    """Страница настроек"""
    # Логируем просмотр настроек
    log_action(current_user.id, current_user.username, 'view_page', 'Просмотр страницы настроек')
    return render_template('settings.html', current_user=current_user)


# ==================== МАРШРУТЫ ДЛЯ ОТЧЕТОВ И ЛОГОВ ====================

@app.route('/reports/devices/excel')
@login_required
def download_excel_report():
    """Скачивание отчета в формате Excel"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('index'))

    try:
        devices = get_all_devices()
        excel_buffer = generate_excel_report(devices)

        # Логируем действие
        log_action(current_user.id, current_user.username, 'download_report',
                   f'Скачан отчет Excel с {len(devices)} устройствами')

        filename = f"devices_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        return send_file(excel_buffer,
                         download_name=filename,
                         as_attachment=True,
                         mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

    except Exception as e:
        logger.error(f"Error generating Excel report: {e}")
        flash('Ошибка генерации отчета Excel', 'error')
        return redirect(url_for('web_devices'))


@app.route('/reports/devices/pdf')
@login_required
def download_pdf_report():
    """Скачивание отчета в формате PDF"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('index'))

    try:
        devices = get_all_devices()
        pdf_buffer = generate_pdf_report(devices)

        # Логируем действие
        log_action(current_user.id, current_user.username, 'download_report',
                   f'Скачан отчет PDF с {len(devices)} устройствами')

        filename = f"devices_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        return send_file(pdf_buffer,
                         download_name=filename,
                         as_attachment=True,
                         mimetype='application/pdf')

    except Exception as e:
        logger.error(f"Error generating PDF report: {e}")
        flash('Ошибка генерации отчета PDF', 'error')
        return redirect(url_for('web_devices'))


@app.route('/reports/devices/docx')
@login_required
def download_docx_report():
    """Скачивание отчета в формате Word"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('index'))

    try:
        devices = get_all_devices()
        docx_buffer = generate_docx_report(devices)

        # Логируем действие
        log_action(current_user.id, current_user.username, 'download_report',
                   f'Скачан отчет Word с {len(devices)} устройствами')

        filename = f"devices_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        return send_file(docx_buffer,
                         download_name=filename,
                         as_attachment=True,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    except Exception as e:
        logger.error(f"Error generating DOCX report: {e}")
        flash('Ошибка генерации отчета Word', 'error')
        return redirect(url_for('web_devices'))


@app.route('/logs')
@login_required
def web_logs():
    """Страница просмотра логов действий"""
    if not current_user.is_admin:
        flash('Доступ запрещен. Требуются права администратора.', 'error')
        return redirect(url_for('index'))

    try:
        logs = get_action_logs(limit=100)
        return render_template('logs.html', logs=logs, current_user=current_user)
    except Exception as e:
        logger.error(f"Error in web_logs route: {e}")
        flash('Ошибка загрузки логов', 'error')
        return render_template('logs.html', logs=[], current_user=current_user)


# ==================== API МАРШРУТЫ ====================

@app.route('/api/auth/login', methods=['POST'])
def api_login():
    """API аутентификация для клиентов"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400

        username = data.get('username')
        password = data.get('password')

        if not username or not password:
            return jsonify({'error': 'Username and password required'}), 400

        user = authenticate_user(username, password)
        if user:
            access_token = create_access_token(identity=user.username)
            return jsonify({
                'access_token': access_token,
                'username': user.username
            })
        else:
            return jsonify({'error': 'Invalid credentials'}), 401

    except Exception as e:
        logger.error(f"Login error: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/secure/submit', methods=['POST'])
@jwt_required()
def secure_submit_data():
    """Безопасный endpoint для отправки данных с клиентов"""
    try:
        current_api_user = get_jwt_identity()
        logger.info(f"Data submission from user: {current_api_user}")

        if not request.data:
            return jsonify({'error': 'No data provided'}), 400

        # Дешифруем данные
        try:
            encrypted_data = request.get_data().decode('utf-8')
            decrypted_data = decrypt_data(encrypted_data)
            data = json.loads(decrypted_data)
        except Exception as e:
            logger.error(f"Decryption error: {e}")
            return jsonify({'error': 'Invalid or corrupted data'}), 400

        # Валидация обязательных полей
        required_fields = ['device_id', 'computer_name', 'mac_address', 'cpu_info']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400

        # Сохраняем данные
        if save_device_data(data):
            return jsonify({
                'status': 'success',
                'message': 'Data received and saved successfully',
                'received_by': current_api_user
            })
        else:
            return jsonify({'error': 'Failed to save data'}), 500

    except Exception as e:
        logger.error(f"Error in secure_submit_data: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/api/devices')
@jwt_required()
def api_devices():
    """API endpoint для получения списка устройств"""
    try:
        devices = get_all_devices()
        return jsonify(devices)
    except Exception as e:
        logger.error(f"Error in api_devices: {e}")
        return jsonify({'error': 'Internal server error'}), 500


@app.route('/health')
def health_check():
    """Проверка здоровья сервера"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT 1')
        conn.close()

        return jsonify({
            'status': 'healthy',
            'database': 'connected',
            'timestamp': datetime.now().isoformat()
        })
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({
            'status': 'unhealthy',
            'database': 'disconnected',
            'error': str(e)
        }), 500

@app.route('/users/change_own_password', methods=['POST'])
@login_required
def change_own_password_route():
    """Смена пароля текущего пользователя с подтверждением старого пароля"""
    try:
        old_password = request.form.get('old_password')
        new_password = request.form.get('new_password')
        confirm_password = request.form.get('confirm_password')

        if not old_password or not new_password or not confirm_password:
            flash('Заполните все поля', 'error')
            return redirect(url_for('web_settings'))

        if new_password != confirm_password:
            flash('Новый пароль и подтверждение не совпадают', 'error')
            return redirect(url_for('web_settings'))

        if len(new_password) < 6:
            flash('Пароль должен содержать минимум 6 символов', 'error')
            return redirect(url_for('web_settings'))

        # Проверяем старый пароль
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT password_hash FROM users WHERE id = ?', (current_user.id,))
        user_data = cursor.fetchone()
        conn.close()

        if not user_data or not bcrypt.check_password_hash(user_data['password_hash'], old_password):
            flash('Неверный старый пароль', 'error')
            return redirect(url_for('web_settings'))

        # Обновляем пароль
        if update_user_password(current_user.id, new_password):
            # Логируем смену пароля
            log_action(current_user.id, current_user.username, 'change_own_password',
                      'Пользователь сменил свой пароль')
            flash('Пароль успешно изменен', 'success')
        else:
            flash('Ошибка изменения пароля', 'error')

        return redirect(url_for('web_settings'))

    except Exception as e:
        logger.error(f"Error changing own password: {e}")
        flash('Ошибка изменения пароля', 'error')
        return redirect(url_for('web_settings'))


if __name__ == '__main__':
    # Создаем папки если их нет
    os.makedirs('templates', exist_ok=True)
    os.makedirs('static', exist_ok=True)

    # Запуск сервера
    print("=" * 60)
    print("🔒 Secure System Information Server with Web Auth")
    print("=" * 60)
    print(f"Server URL: {Config.SERVER_URL}")
    print(f"Web Login: {Config.SERVER_URL}/login")
    print("Default admin credentials: admin / admin123")
    print("=" * 60)

    try:
        app.run(
            host=Config.SERVER_HOST,
            port=int(Config.SERVER_PORT),
            debug=True
        )
    except Exception as e:
        logger.error(f"Server failed to start: {e}")